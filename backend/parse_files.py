import pypdf
import openpyxl
from docx import Document as DocxDocument
from langchain_classic.schema import Document
from io import BytesIO

def parse_file(file_bytes: bytes, file_name: str, s3_key: str) -> list[Document]:
    ext = file_name.rsplit(".", 1)[-1].lower()

    if ext == "pdf":
        reader = pypdf.PdfReader(BytesIO(file_bytes))
        return [
            Document(
                page_content=page.extract_text(),
                metadata={
                    "s3_key": s3_key,
                    "file_name": file_name,
                    "file_type": "pdf",
                    "page": i,
                }
            )
            for i, page in enumerate(reader.pages)
            if page.extract_text().strip()
        ]

    elif ext in ("xlsx", "xls", "csv"):
        wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True)
        docs = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []
            headers = []
            for i, row in enumerate(sheet.iter_rows(values_only=True)):
                row_text = [str(cell) if cell is not None else "" for cell in row]
                if i == 0:
                    headers = row_text
                else:
                    rows.append(" | ".join(row_text))

            if rows:
                content = f"Headers: {' | '.join(headers)}\n" + "\n".join(rows)
                docs.append(Document(
                    page_content=content,
                    metadata={
                        "s3_key": s3_key,
                        "file_name": file_name,
                        "file_type": "excel",
                        "sheet": sheet_name,
                        "row_count": len(rows),
                    }
                ))
        return docs

    elif ext in ("docx", "doc"):
        doc = DocxDocument(BytesIO(file_bytes))
        docs = []
        current_section = "Introduction"
        current_text = []

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                # Save previous section
                if current_text:
                    docs.append(Document(
                        page_content="\n".join(current_text),
                        metadata={
                            "s3_key": s3_key,
                            "file_name": file_name,
                            "file_type": "docx",
                            "section": current_section,
                        }
                    ))
                current_section = para.text
                current_text = []
            else:
                if para.text.strip():
                    current_text.append(para.text)

        # Save last section
        if current_text:
            docs.append(Document(
                page_content="\n".join(current_text),
                metadata={
                    "s3_key": s3_key,
                    "file_name": file_name,
                    "file_type": "docx",
                    "section": current_section,
                }
            ))
        return docs

    elif ext in ("png", "jpg", "jpeg", "webp"):
        # Use Azure AI Vision or Tesseract for OCR
        # For now, placeholder
        return [Document(
            page_content="[Image — OCR not yet implemented]",
            metadata={
                "s3_key": s3_key,
                "file_name": file_name,
                "file_type": "image",
            }
        )]

    else:
        # Plain text fallback
        return [Document(
            page_content=file_bytes.decode("utf-8", errors="ignore"),
            metadata={
                "s3_key": s3_key,
                "file_name": file_name,
                "file_type": ext,
            }
        )]